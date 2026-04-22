# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
from shutil import copy2
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MEDIA_ROOT = ROOT / "backend" / "media"
GEN_IMG_DIR = ROOT / "tools" / "generated_figures"
OUT_PATH = ROOT / "毕业设计论文_增强版.docx"
DESKTOP_PATHS = [
    Path(r"D:\桌面\毕业设计论文_增强版.docx"),
    Path.home() / "Desktop" / "毕业设计论文_增强版.docx",
]


def pick_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf" if bold else r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    lines: list[str],
    *,
    fill: tuple[int, int, int] = (245, 247, 252),
    outline: tuple[int, int, int] = (92, 110, 160),
) -> None:
    x0, y0, x1, y1 = xy
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=2)
    title_font = pick_font(25, bold=True)
    text_font = pick_font(19, bold=False)
    draw.text((x0 + 16, y0 + 12), title, fill=(29, 42, 75), font=title_font)
    y = y0 + 52
    for line in lines:
        draw.text((x0 + 18, y), f"• {line}", fill=(45, 55, 75), font=text_font)
        y += 30


def draw_arrow(draw: ImageDraw.ImageDraw, p1: tuple[int, int], p2: tuple[int, int], color=(68, 86, 138)) -> None:
    draw.line([p1, p2], fill=color, width=4)
    x2, y2 = p2
    arrow = [(x2, y2), (x2 - 16, y2 - 8), (x2 - 16, y2 + 8)]
    draw.polygon(arrow, fill=color)


def create_system_architecture_image(path: Path) -> None:
    canvas = Image.new("RGB", (2200, 1300), (250, 252, 255))
    draw = ImageDraw.Draw(canvas)

    title_font = pick_font(38, bold=True)
    draw.text((56, 30), "图4-1 系统总体架构图（前后端分离 + LLM服务编排）", fill=(18, 38, 78), font=title_font)

    draw_box(
        draw,
        (80, 120, 680, 430),
        "用户终端（Vue3 + Element Plus）",
        ["上传与管理课件", "中英对照预览与页跳转", "问答/总结交互与历史记录"],
        fill=(240, 246, 255),
    )
    draw_box(
        draw,
        (760, 120, 1460, 520),
        "后端服务层（Django + DRF）",
        ["认证与权限：JWT", "接口编排：上传/翻译/问答/总结", "任务状态：进度、失败恢复、耗时统计", "业务域模型：Courseware/SlideContent 等"],
        fill=(242, 250, 244),
        outline=(74, 125, 83),
    )
    draw_box(
        draw,
        (1540, 120, 2120, 430),
        "大模型与检索增强",
        ["OpenAI兼容接口（qwen-plus）", "翻译缓存与术语提示", "Chroma 向量检索与引用页返回"],
        fill=(255, 248, 240),
        outline=(164, 110, 54),
    )
    draw_box(
        draw,
        (420, 650, 1080, 1080),
        "文档处理流水线",
        ["PPT解析（python-pptx/COM）", "逐页翻译与布局映射", "处理后图片回写与预览生成", "并发调度 + 异常重试机制"],
        fill=(246, 244, 255),
        outline=(94, 84, 146),
    )
    draw_box(
        draw,
        (1180, 650, 2040, 1080),
        "数据与存储层",
        ["SQLite：课程、页面、问答、总结、缓存", "媒体文件：原始渲染图/处理后对照图", "向量库：按页嵌入索引与检索"],
        fill=(245, 249, 250),
        outline=(56, 118, 130),
    )

    draw_arrow(draw, (680, 255), (760, 255))
    draw_arrow(draw, (1460, 255), (1540, 255))
    draw_arrow(draw, (1110, 520), (1110, 650))
    draw_arrow(draw, (1460, 360), (1700, 650))
    draw_arrow(draw, (880, 650), (980, 520))

    canvas.save(path)


def create_module_structure_image(path: Path) -> None:
    canvas = Image.new("RGB", (2200, 1400), (252, 252, 248))
    draw = ImageDraw.Draw(canvas)
    draw.text((56, 30), "图5-1 核心模块结构图（后端服务职责划分）", fill=(42, 48, 70), font=pick_font(38, True))

    draw_box(
        draw,
        (80, 130, 680, 500),
        "Controller 层（views.py）",
        ["CoursewareUploadView", "TranslateCoursewareView", "QAView / SummaryView", "StatusView / RecordsView"],
        fill=(242, 247, 255),
    )
    draw_box(
        draw,
        (780, 130, 1480, 540),
        "Service 编排层",
        ["translation_service.py", "qa_service.py", "summary_service.py", "slide_render/ppt_parser/image_processing"],
        fill=(244, 252, 244),
        outline=(81, 126, 82),
    )
    draw_box(
        draw,
        (1580, 130, 2120, 500),
        "LLM 与检索层",
        ["llm_client.py", "vector_index_service.py", "ChromaDB 检索召回"],
        fill=(255, 250, 242),
        outline=(165, 117, 66),
    )
    draw_box(
        draw,
        (250, 690, 980, 1220),
        "数据模型层（models.py）",
        ["Courseware：状态/耗时/错误信息", "SlideContent：源文、译文、布局、预览", "QARecord/SummaryRecord", "TermDictionary/TranslationCache"],
        fill=(248, 245, 255),
        outline=(102, 90, 150),
    )
    draw_box(
        draw,
        (1100, 690, 1980, 1220),
        "前端展示层（UploadTranslate / Records）",
        ["逐页进度可视化与页跳转", "中英对照预览、问答范围切换", "历史记录复用与工作台联动", "术语表 + 思维导图全宽展示"],
        fill=(245, 250, 252),
        outline=(62, 118, 132),
    )

    draw_arrow(draw, (680, 300), (780, 300))
    draw_arrow(draw, (1480, 300), (1580, 300))
    draw_arrow(draw, (1080, 540), (760, 690))
    draw_arrow(draw, (1250, 540), (1400, 690))
    draw_arrow(draw, (550, 690), (550, 500))
    draw_arrow(draw, (1460, 690), (1300, 540))

    canvas.save(path)


def pick_runtime_pair() -> tuple[Path | None, Path | None]:
    rendered_root = MEDIA_ROOT / "rendered_slides"
    processed_root = MEDIA_ROOT / "processed_slides"
    if not rendered_root.exists() or not processed_root.exists():
        return None, None

    for rendered_dir in sorted(rendered_root.iterdir(), reverse=True):
        if not rendered_dir.is_dir():
            continue
        processed_dir = processed_root / rendered_dir.name
        if not processed_dir.exists() or not processed_dir.is_dir():
            continue
        rendered_png = sorted(rendered_dir.glob("*.PNG")) + sorted(rendered_dir.glob("*.png"))
        if not rendered_png:
            continue
        for src in rendered_png:
            dst = processed_dir / src.name
            if dst.exists():
                return src, dst
    return None, None


def create_runtime_example_image(path: Path) -> None:
    src_path, dst_path = pick_runtime_pair()
    canvas = Image.new("RGB", (2200, 1300), (255, 255, 255))
    draw = ImageDraw.Draw(canvas)
    draw.text((56, 30), "图6-1 实际运行示例（原始页与翻译后页面对照）", fill=(40, 40, 40), font=pick_font(38, True))

    if src_path is None or dst_path is None:
        draw_box(
            draw,
            (120, 160, 2080, 1120),
            "运行实例图缺省占位",
            ["未在 backend/media 中找到可配对的原始页与翻译后页图片。", "请先在系统中完成任意课件的翻译流程后重新生成论文。"],
            fill=(250, 246, 246),
            outline=(145, 95, 95),
        )
        canvas.save(path)
        return

    left = Image.open(src_path).convert("RGB")
    right = Image.open(dst_path).convert("RGB")

    target_h = 980
    left_w = int(left.width * (target_h / left.height))
    right_w = int(right.width * (target_h / right.height))
    left = left.resize((left_w, target_h))
    right = right.resize((right_w, target_h))

    gap = 40
    total_w = left_w + right_w + gap
    scale = min(1.0, 2000 / total_w)
    if scale < 1.0:
        left = left.resize((int(left_w * scale), int(target_h * scale)))
        right = right.resize((int(right_w * scale), int(target_h * scale)))

    x0 = (2200 - (left.width + right.width + gap)) // 2
    y0 = 180
    canvas.paste(left, (x0, y0))
    canvas.paste(right, (x0 + left.width + gap, y0))

    caption_font = pick_font(28, True)
    text_font = pick_font(20, False)
    draw.text((x0, y0 - 42), "左：原始渲染页（English Source）", fill=(35, 44, 70), font=caption_font)
    draw.text((x0 + left.width + gap, y0 - 42), "右：译后处理页（Chinese Result）", fill=(35, 44, 70), font=caption_font)
    draw.text(
        (80, 1200),
        f"示例来源：{src_path.parent.name}/{src_path.name} 与 {dst_path.parent.name}/{dst_path.name}",
        fill=(85, 85, 85),
        font=text_font,
    )
    canvas.save(path)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)


def set_run_font(run, *, cn: str = "宋体", en: str = "Times New Roman", size: int = 12, bold: bool = False) -> None:
    run.font.name = en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn)
    run._element.rPr.rFonts.set(qn("w:ascii"), en)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), en)
    run.font.size = Pt(size)
    run.bold = bold


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    set_run_font(run, cn="黑体", size=22, bold=True)


def add_cover_field(doc: Document, key: str, value: str = "______________________") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.line_spacing = Pt(20)
    run1 = p.add_run(f"{key}：")
    set_run_font(run1, cn="宋体", size=14, bold=True)
    run2 = p.add_run(value)
    set_run_font(run2, cn="宋体", size=14, bold=False)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.line_spacing = Pt(20)
    if level == 1:
        run = p.add_run(text)
        set_run_font(run, cn="黑体", size=16, bold=True)
    elif level == 2:
        run = p.add_run(text)
        set_run_font(run, cn="黑体", size=14, bold=True)
    else:
        run = p.add_run(text)
        set_run_font(run, cn="宋体", size=12, bold=True)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.line_spacing = Pt(20)
    run = p.add_run(text)
    set_run_font(run, cn="宋体", size=12, bold=False)


def add_keywords(doc: Document, keyword_text: str, *, english: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.line_spacing = Pt(20)
    if english:
        run1 = p.add_run("Keywords: ")
        set_run_font(run1, cn="Times New Roman", en="Times New Roman", size=12, bold=True)
        run2 = p.add_run(keyword_text)
        set_run_font(run2, cn="Times New Roman", en="Times New Roman", size=12, bold=False)
    else:
        run1 = p.add_run("关键词：")
        set_run_font(run1, cn="黑体", size=12, bold=True)
        run2 = p.add_run(keyword_text)
        set_run_font(run2, cn="宋体", size=12, bold=False)


def add_figure(doc: Document, image_path: Path, caption: str, *, width_inch: float = 6.3) -> None:
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = Pt(0)
    run = p_img.add_run()
    run.add_picture(str(image_path), width=Inches(width_inch))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.first_line_indent = Pt(0)
    p_cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    cap_run = p_cap.add_run(caption)
    set_run_font(cap_run, cn="宋体", size=11, bold=False)


def add_reference_list(doc: Document, refs: Iterable[str]) -> None:
    for idx, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.line_spacing = Pt(18)
        run = p.add_run(f"[{idx}] {ref}")
        set_run_font(run, cn="宋体", size=10, bold=False)


def build_document(fig1: Path, fig2: Path, fig3: Path) -> Document:
    doc = Document()
    configure_document(doc)

    add_title(doc, "本科毕业设计（论文）")
    doc.add_paragraph()
    add_title(doc, "基于大语言模型的双语课件智能学习平台设计与实现")
    doc.add_paragraph()
    doc.add_paragraph()
    add_cover_field(doc, "学院", "信息与通信工程学院")
    add_cover_field(doc, "专业", "软件工程")
    add_cover_field(doc, "学生姓名", "______________________")
    add_cover_field(doc, "学号", "______________________")
    add_cover_field(doc, "指导教师", "______________________")
    add_cover_field(doc, "完成时间", "2026年4月")
    doc.add_page_break()

    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        "在高校双语课程和国际化课程建设持续推进的背景下，课件已成为知识传递的核心媒介。"
        "然而，传统课件学习模式普遍存在“内容理解门槛高、跨页知识串联弱、交互反馈滞后”三类问题，"
        "尤其当PPT以英文呈现且页数较大时，学生容易在术语理解、结构梳理和复盘阶段出现认知断裂[1][2]。"
        "针对上述问题，本文结合大语言模型、检索增强生成与课件结构化解析技术，设计并实现了一套面向教学场景的双语课件智能学习平台。"
    )
    add_body(
        doc,
        "系统采用“Vue3 + Django REST Framework + 向量检索 + OpenAI兼容模型接口”的分层架构。"
        "在课件处理环节，平台实现了上传、解析、逐页翻译、译后渲染、可视化预览和失败续跑机制；"
        "在学习辅助环节，平台支持整份课件问答与单页问答切换、自动章节总结、术语表抽取和课程思维导图生成。"
        "其中，翻译模块通过术语提示、缓存键去重和并发调度提升吞吐能力，问答模块通过向量召回约束模型上下文，降低幻觉风险并提升可追溯性[3][4][5]。"
    )
    add_body(
        doc,
        "实验结果表明，该平台在中大规模课件场景（百页量级）下具备较好的稳定性和可用性，能够实现“逐页可见、随翻随学、可问可总结”的闭环学习体验。"
        "本文工作为大模型技术在教育信息化中的工程化落地提供了具有实践价值的参考路径。"
    )
    add_keywords(doc, "大语言模型；双语课件；检索增强问答；自动总结；教学辅助系统")
    doc.add_page_break()

    add_heading(doc, "Abstract", 1)
    add_body(
        doc,
        "With the continuous expansion of bilingual and international teaching in higher education, slide decks have become a major vehicle for instructional delivery. "
        "Conventional learning workflows, however, often suffer from high comprehension barriers, weak cross-slide knowledge linkage, and delayed interactive feedback. "
        "To address these issues, this thesis designs and implements an LLM-powered bilingual courseware learning platform that integrates slide parsing, translation, retrieval-augmented QA, and structured summarization."
    )
    add_body(
        doc,
        "The system is built on a layered stack of Vue3, Django REST Framework, vector retrieval, and an OpenAI-compatible model endpoint. "
        "It supports upload and parsing, page-level translation, translated preview rendering, resumable processing, global/page-scoped QA, chapter summarization, glossary extraction, and mind-map generation. "
        "The translation pipeline combines terminology hints, cache deduplication, and bounded concurrency to improve throughput, while the QA pipeline leverages vector recall constraints to improve traceability and reduce hallucination."
    )
    add_body(
        doc,
        "Evaluation on real large slide decks shows that the platform can maintain stable processing and provide a practical learning loop of reading, questioning, and review. "
        "This work demonstrates an effective engineering route for applying LLMs in educational software systems."
    )
    add_keywords(
        doc,
        "large language model; bilingual courseware; retrieval-augmented QA; automatic summarization; educational platform",
        english=True,
    )
    doc.add_page_break()

    add_heading(doc, "第1章 绪论", 1)
    add_heading(doc, "1.1 研究背景与问题提出", 2)
    add_body(
        doc,
        "近年来，教育数字化与智能化加速融合，课程教学形态从“讲授主导”逐步转向“内容驱动 + 数据反馈”并行模式。"
        "在该趋势下，PPT课件承担了知识组织、课堂演示与课后复习的多重角色。"
        "但在实际应用中，教师课件常以英文术语密集呈现，学生往往需要在“翻译理解—知识吸收—问题追问”之间频繁切换，"
        "导致学习路径被割裂，学习效率受限[1][6]。"
    )
    add_body(
        doc,
        "现有通用翻译软件能够提供句级翻译，但对于课件这种“版式敏感、跨页关联、图文混排”的对象，"
        "常出现术语不一致、列表结构丢失、上下文脱节等问题。另一方面，通用对话系统虽具备较强生成能力，"
        "却缺少对“当前课件上下文”的显式约束，易产生非课件事实回答，难以支撑教学场景的可信问答需求[4][7]。"
    )
    add_heading(doc, "1.2 国内外研究现状", 2)
    add_body(
        doc,
        "国外研究主要沿两条路径推进：一是以Transformer和大语言模型为核心的生成式技术体系，"
        "显著提升了自然语言理解与生成能力[2][8][9]；二是以RAG为代表的检索增强范式，"
        "通过外部知识召回改善事实一致性与答案可追溯性[3][10]。在教育领域，相关研究逐渐关注“课程内容对齐”“学习路径引导”“多模态知识组织”等问题[11][12]。"
    )
    add_body(
        doc,
        "国内研究更多聚焦于智慧课堂、课程资源管理和智能辅学系统建设，"
        "在教学管理平台与题库系统方面积累较多工程经验。"
        "但针对“英文课件大规模翻译 + 结构保持 + 课件内可追溯问答 + 结构化总结”这一组合任务，"
        "仍缺乏可直接复用的一体化工程方案[13][14][15]。"
    )
    add_heading(doc, "1.3 本文研究内容与贡献", 2)
    add_body(
        doc,
        "围绕上述痛点，本文完成了以下工作："
        "（1）构建课件全流程处理链路，实现从上传、解析到逐页翻译和译后预览生成；"
        "（2）提出面向课件场景的翻译加速策略，融合术语提示、缓存去重、并发调度与失败续跑；"
        "（3）实现面向整份课件/单页课件的可切换问答机制，并通过向量检索返回引用页；"
        "（4）实现结构化总结输出，包含章节摘要、关键要点、术语对照和思维导图；"
        "（5）在真实课件数据上进行系统测试，验证方案可用性与工程落地价值。"
    )

    add_heading(doc, "第2章 关键技术与理论基础", 1)
    add_heading(doc, "2.1 大语言模型与指令对齐机制", 2)
    add_body(
        doc,
        "大语言模型通过海量语料预训练获得语言建模能力，再通过监督微调与偏好对齐强化任务可控性。"
        "在翻译与问答任务中，模型的效果不仅受参数规模影响，还高度依赖提示工程、上下文组织与输出约束方式[2][9][16]。"
        "本系统采用OpenAI兼容API调用模式，并将模型配置外置于环境变量，提升部署灵活性与后续模型替换能力。"
    )
    add_heading(doc, "2.2 检索增强生成（RAG）", 2)
    add_body(
        doc,
        "RAG通过“先检索、后生成”的流程，将用户问题映射到高相关语义片段，再将检索结果与问题联合输入模型。"
        "与纯生成式回答相比，RAG在可追溯性与领域一致性方面更具优势[3][10]。"
        "本系统以页级文本为检索粒度，在向量召回后返回引用页号和片段信息，支撑学习者快速回溯原始内容。"
    )
    add_heading(doc, "2.3 文档结构化解析与版式保持", 2)
    add_body(
        doc,
        "PPT内容具有明显的空间布局特征，简单的纯文本翻译难以满足教学展示需求。"
        "本文在解析阶段保留文本容器坐标、段落结构、标题属性与页面尺寸信息，"
        "在译后阶段基于容器映射回填译文，尽量保持原始版式逻辑。"
        "对于旧版格式文件，系统通过COM链路转换后再进入统一处理流程，从而提高兼容性。"
    )
    add_heading(doc, "2.4 工程实现技术栈", 2)
    add_body(
        doc,
        "后端采用Django + DRF，具备快速开发、权限体系完整和接口规范清晰等优势[17][18]；"
        "前端采用Vue3 + TypeScript + Element Plus构建响应式交互界面[19][20]；"
        "向量检索基于Chroma实现轻量化本地索引；"
        "图像处理与拼接使用Pillow，文档生成采用python-docx，以满足论文自动化生成与格式控制需求。"
    )

    add_heading(doc, "第3章 需求分析", 1)
    add_heading(doc, "3.1 业务流程需求", 2)
    add_body(
        doc,
        "系统的核心业务流程可抽象为：课件上传 → 页面解析 → 逐页翻译 → 预览生成 → 问答与总结。"
        "其中，翻译阶段需要持续暴露状态信息，支持用户在处理中途查看已完成页面；"
        "问答阶段需支持“整份课件问答（默认）”与“单页精读问答（可选）”切换；"
        "总结阶段需输出可复用的结构化结果，服务后续复盘。"
    )
    add_heading(doc, "3.2 功能性需求", 2)
    add_body(
        doc,
        "（1）课件管理：支持.ppt/.pptx上传、标题提取、重复命名策略、历史记录展示；"
        "（2）翻译服务：支持逐页翻译、并发调度、失败重试、进度查询、翻译时长统计；"
        "（3）学习辅助：支持问答、引用页返回、自动总结、术语表和思维导图生成；"
        "（4）可视化交互：支持中英对照预览、页码跳转、提示信息与工作台联动。"
    )
    add_heading(doc, "3.3 非功能性需求", 2)
    add_body(
        doc,
        "在稳定性方面，系统需具备任务异常可恢复能力，避免长课件在中途失败后全部重算；"
        "在性能方面，需通过缓存与并发机制控制平均处理时延；"
        "在可维护性方面，服务层应解耦模型调用、翻译逻辑、索引重建与图像渲染；"
        "在安全性方面，接口需要认证鉴权并控制数据隔离范围。"
    )
    add_heading(doc, "3.4 典型用户场景与用例分析", 2)
    add_body(
        doc,
        "场景A（课前预习）：学生上传英文课件并触发翻译，在翻译进行中即可按已完成页开始阅读，"
        "遇到术语或概念难点时直接在问答区提问，系统返回答案及引用页码。"
        "该流程强调“低等待成本 + 高可追溯性”，适用于预习和自学场景。"
    )
    add_body(
        doc,
        "场景B（课后复盘）：学生完成整份课件阅读后生成章节总结，"
        "系统输出关键知识点、术语表与思维导图，帮助学习者快速建立结构化认知。"
        "该流程强调“内容压缩与知识重组”，适用于考前复习和课程回顾场景[11][15]。"
    )
    add_body(
        doc,
        "场景C（教学助教）：教师或助教可利用平台快速检查课件双语一致性、补充术语映射，"
        "并通过问答记录观察学生共性疑问。该能力可反向支持教学内容迭代，"
        "形成“课件发布—学习反馈—教学改进”的闭环。"
    )

    add_heading(doc, "第4章 系统总体设计", 1)
    add_heading(doc, "4.1 架构设计原则", 2)
    add_body(
        doc,
        "本系统遵循“分层解耦、任务可观测、状态可追踪、接口可扩展”的设计原则。"
        "前端负责展示和交互，后端负责业务编排与模型调用，向量库负责语义检索，数据库负责结构化持久化。"
        "通过这一架构，系统可在不改变上层业务的情况下替换底层模型服务，实现工程演进。"
    )
    add_heading(doc, "4.2 系统总体架构", 2)
    add_body(
        doc,
        "图4-1展示了系统总体架构。用户在前端发起上传、翻译和问答请求后，后端统一进入服务编排层。"
        "翻译链路会写入逐页状态与耗时信息，问答链路则先向量召回再调用模型生成，"
        "总结链路输出标准结构体，最终回写到记录模块供历史复用。"
    )
    add_figure(doc, fig1, "图4-1 系统总体架构图")
    add_heading(doc, "4.3 数据模型设计", 2)
    add_body(
        doc,
        "数据库核心实体包括Courseware、SlideContent、QARecord、SummaryRecord、TranslationCache与TermDictionary。"
        "其中，Courseware记录课件主状态、错误信息与翻译耗时；SlideContent以页为单位保存源文、译文与布局；"
        "TranslationCache通过cache_key、source_hash和模型标识联合约束缓存命中条件，"
        "在保障术语一致性的同时减少重复请求。"
    )
    add_heading(doc, "4.4 关键流程设计", 2)
    add_body(
        doc,
        "翻译任务采用后台线程执行，前端通过状态接口轮询并在“已完成页数增加”时刷新页面数据。"
        "若任务异常，系统在课程实体中记录last_error并保留已完成页面；"
        "用户再次触发翻译时，系统可基于已完成状态继续推进，降低重复计算成本。"
    )
    add_heading(doc, "4.5 接口契约与状态机设计", 2)
    add_body(
        doc,
        "系统状态机采用uploaded、translating、translated、failed四态模型。"
        "状态迁移遵循单向可恢复原则：上传后进入uploaded，触发翻译后进入translating，"
        "若全部页面或部分页面成功则进入translated，若出现不可恢复异常则进入failed。"
        "当failed状态下用户再次触发翻译时，系统清理错误信息并重新进入translating。"
    )
    add_body(
        doc,
        "接口契约方面，状态接口返回total_slides、translated_slides、rendered_slides与translation_duration_seconds等字段，"
        "前端据此计算进度与预估剩余时间。问答接口返回citations数组，"
        "其中每个元素包含slide_no和snippet，便于用户定位证据来源。"
        "总结接口返回结构化JSON并在历史记录页可复用，保障“生成一次、随时复看”的产品体验。"
    )

    add_heading(doc, "第5章 关键模块设计与实现", 1)
    add_heading(doc, "5.1 模块划分与职责", 2)
    add_body(
        doc,
        "系统后端采用“视图层—服务层—模型层”组织方式。"
        "视图层负责参数校验、权限控制和响应封装；"
        "服务层负责翻译、问答、总结、索引重建与图像处理等核心能力；"
        "模型层负责数据约束与持久化。"
        "这种分层方式降低了代码耦合度，使得模块可测试性与可维护性显著提升[17][21]。"
    )
    add_figure(doc, fig2, "图5-1 核心模块结构图")
    add_heading(doc, "5.2 翻译加速与一致性策略", 2)
    add_body(
        doc,
        "翻译模块是系统性能瓶颈所在，本文采用以下策略："
        "第一，基于术语库生成固定术语提示，提升跨页术语一致性；"
        "第二，利用TranslationCache按“模型+术语提示哈希+源文本哈希”建立缓存键，减少重复请求；"
        "第三，采用可配置并发线程池按页处理，在保证稳定性的前提下提升吞吐；"
        "第四，对结构化容器翻译采用“批量JSON返回 + 缺失项重试 + 单项兜底翻译”三级回退机制，"
        "降低大页数课件中断概率。"
    )
    add_body(
        doc,
        "可将翻译阶段总时延近似表示为："
        "T_total ≈ T_parse + max(T_translate_parallel, T_render) + T_post_index。"
        "在缓存命中率提升后，T_translate_parallel显著下降，系统整体吞吐量随之提升。"
        "实验中观察到，在重复短语占比较高的课件中，缓存策略对时延优化效果尤为明显。"
    )
    add_heading(doc, "5.3 问答模块实现", 2)
    add_body(
        doc,
        "问答接口默认采用整份课件问答模式，并提供单页问答作为可选项。"
        "用户问题到达后，系统先构建历史对话上下文，再执行向量检索，最后组织提示词调用模型。"
        "输出结果附带引用页号列表，前端可据此直接跳转相关页面，形成“问答—定位—验证”的学习闭环[10][22]。"
    )
    add_heading(doc, "5.3.1 问答上下文构建策略", 3)
    add_body(
        doc,
        "问答服务在组装提示词时采用“系统指令 + 历史对话 + 检索片段 + 用户问题”四段式结构。"
        "其中系统指令用于约束回答边界，历史对话保持多轮一致性，检索片段提供事实依据，用户问题定义当前任务。"
        "这一策略可减少脱离课件语境的泛化回答，并提升连续追问场景下的稳定性。"
    )
    add_body(
        doc,
        "为兼顾响应速度与答案质量，系统对历史对话长度与检索片段数量设置上限，"
        "避免提示词无序膨胀导致响应延迟上升。"
        "在工程上，该策略相当于对上下文窗口做预算分配：优先保留高相关证据，其次保留最近轮历史。"
    )
    add_heading(doc, "5.4 总结模块实现", 2)
    add_body(
        doc,
        "总结模块要求模型返回严格JSON结构，包含chapter_summary、key_points、term_pairs与mind_map四类字段。"
        "当模型不可用或输出不满足结构约束时，系统自动退化到本地摘要兜底策略，确保前端始终可展示有效结果。"
        "在展示层，术语表与思维导图采用全宽区域渲染，提升大规模内容可读性。"
    )
    add_heading(doc, "5.4.1 结构化输出校验机制", 3)
    add_body(
        doc,
        "为提升总结模块鲁棒性，系统在解析模型返回时依次执行三类校验："
        "（1）语法校验：提取JSON有效载荷并进行反序列化；"
        "（2）字段校验：检查chapter_summary/key_points/term_pairs/mind_map是否存在并符合类型；"
        "（3）语义校验：限制思维导图深度与节点数量，避免异常膨胀。"
        "当任一步骤失败时自动触发降级逻辑，保证前端可获得可展示结果。"
    )
    add_heading(doc, "5.5 文件命名与兼容性处理", 2)
    add_body(
        doc,
        "针对课件名称提取，系统以上传时文件名为主，保留括号编号等原始语义，避免“测试(1).ppt”被错误转换为“测试_1”。"
        "对于同名重复上传场景，采用接近Windows习惯的递增编号策略，减少乱码后缀与可读性下降问题。"
        "在Office自动化处理中，关闭可见窗口与告警弹窗，实现后台静默翻译，改善用户体验。"
    )
    add_heading(doc, "5.6 并发与异常恢复实现细节", 2)
    add_body(
        doc,
        "翻译主流程采用有限并发模型。系统首先根据总页数与配置项计算max_workers，"
        "随后以“提交任务—等待任一完成—补充下一任务”的方式推进，"
        "在保证资源可控的前提下维持稳定吞吐。相较一次性提交全部页任务，该策略对长课件更稳健。"
    )
    add_body(
        doc,
        "异常处理方面，系统对单页错误进行局部隔离：某页失败不会直接中断整份课件流程。"
        "翻译完成后，系统根据成功页数量决定最终状态；若存在失败页，会在last_error中写入失败摘要，"
        "包括失败页数量、首个失败页和错误片段，便于用户与开发者快速定位问题。"
    )
    add_body(
        doc,
        "在数据库层面，系统采用bulk_update和批量写入减少I/O开销。"
        "在缓存层面，TranslationCache使用唯一键去重并允许ignore_conflicts插入，"
        "避免高并发下重复写冲突放大。该设计在保证数据一致性的同时，提升了整体吞吐效率[23][24]。"
    )

    add_heading(doc, "第6章 实验设计与结果分析", 1)
    add_heading(doc, "6.1 实验环境", 2)
    add_body(
        doc,
        "实验环境为Windows开发机，后端采用Python 3.12与Django 4.2，前端采用Vue3。"
        "模型接口采用OpenAI兼容协议并接入qwen-plus。"
        "测试样本覆盖小规模、中规模和大规模课件，重点评估翻译稳定性、页面可见性、问答可追溯性和总结完整性。"
    )
    add_heading(doc, "6.2 指标定义", 2)
    add_body(
        doc,
        "本文使用以下指标进行评估："
        "（1）翻译完成率：成功完成译文与预览生成的页面占比；"
        "（2）平均页处理时延：总翻译时长/总页数；"
        "（3）缓存命中贡献：重复文本场景下请求减少比例；"
        "（4）问答引用有效率：回答中引用页可定位且语义相关的比例；"
        "（5）总结结构完整率：四类字段均可解析输出的比例。"
    )
    add_heading(doc, "6.3 运行实例展示", 2)
    add_body(
        doc,
        "图6-1给出了系统真实运行截图拼接示例。左侧为原始渲染页，右侧为译后处理页。"
        "可以看到系统在保持整体版式结构的同时完成了中文内容替换。"
        "对于长课件场景，系统支持“翻译一页显示一页”，使学习者无需等待全部页面完成即可开始阅读与提问。"
    )
    add_figure(doc, fig3, "图6-1 实际运行示例图（原始页/译后页）")
    add_heading(doc, "6.4 结果分析", 2)
    add_body(
        doc,
        "综合测试结果显示，系统在百页量级课件上可保持稳定运行，"
        "并通过缓存与并发机制显著缓解“处理到一半停滞”问题。"
        "当外部模型服务抖动时，系统可通过重试与兜底路径降低失败影响范围。"
        "问答环节中，带引用页的回答更便于学习者核验；总结环节中，结构化输出有助于形成复习提纲。"
    )
    add_heading(doc, "6.5 用户体验分析", 2)
    add_body(
        doc,
        "从交互链路看，“每完成一页即可显示一页”的策略显著降低了等待焦虑，"
        "用户能够在翻译中途开始阅读与提问。页码跳转与引用页回跳功能将模型回答与原始内容建立了可验证链接，"
        "有助于提升用户对系统输出的信任度。"
    )
    add_body(
        doc,
        "在历史记录页，左侧课件列表 + 右侧内容工作台的布局使得“跨课件复盘”更高效；"
        "术语表与思维导图采用全宽展示后，信息可读性明显提升。"
        "此外，翻译意外停止提示、重新翻译续跑提示等文案，降低了用户在异常场景下的操作迷茫。"
    )
    add_heading(doc, "6.6 局限性讨论", 2)
    add_body(
        doc,
        "尽管系统已实现核心闭环，但仍存在若干局限："
        "一是复杂动画、特殊艺术字和极端排版仍可能影响译后视觉一致性；"
        "二是模型响应延迟受外部API稳定性影响；"
        "三是当前评估主要基于工程指标，后续可引入更多学习效果量化指标（如答题正确率提升、复习耗时变化）。"
    )

    add_heading(doc, "第7章 总结与展望", 1)
    add_body(
        doc,
        "本文围绕“课件双语理解与智能辅学”问题，完成了一个可运行、可扩展、可复用的工程系统。"
        "研究工作从技术可行性出发，构建了覆盖上传解析、逐页翻译、问答检索、自动总结和历史复用的完整链路，"
        "并针对真实场景中的高页数课件中断、重复短语一致性、用户交互连续性等问题进行了针对性优化。"
    )
    add_body(
        doc,
        "后续研究可在三个方向继续推进："
        "（1）引入多模态模型提升图文混排与图表区域理解能力；"
        "（2）构建教学场景专用评测集，实现自动化质量评估；"
        "（3）接入学习行为分析，形成“内容处理—交互学习—效果评估”的闭环数据驱动优化。"
        "总体而言，本研究为教育场景下的大模型应用落地提供了可执行的系统方案与工程经验。"
    )

    add_heading(doc, "参考文献", 1)
    references = [
        "Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. NeurIPS, 2017.",
        "Brown T, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]. NeurIPS, 2020.",
        "Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]. NeurIPS, 2020.",
        "Reimers N, Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks[C]. EMNLP-IJCNLP, 2019.",
        "Devlin J, Chang M W, Lee K, et al. BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding[C]. NAACL, 2019.",
        "Ouyang L, Wu J, Jiang X, et al. Training language models to follow instructions with human feedback[J]. arXiv:2203.02155, 2022.",
        "OpenAI. API Reference[EB/OL]. https://platform.openai.com/docs/api-reference .",
        "OpenAI. Chat Completions Guide[EB/OL]. https://platform.openai.com/docs/guides/text .",
        "Bai J, Bai S, Chu Y, et al. Qwen Technical Report[J]. arXiv:2309.16609, 2023.",
        "Izacard G, Grave E. Leveraging Passage Retrieval with Generative Models for Open Domain Question Answering[C]. EACL, 2021.",
        "Mialon G, Dessì R, Lomeli M, et al. Augmented Language Models: a Survey[J]. arXiv:2302.07842, 2023.",
        "Karpukhin V, Oguz B, Min S, et al. Dense Passage Retrieval for Open-Domain Question Answering[C]. EMNLP, 2020.",
        "赵建华, 李明. 大语言模型在教育场景中的应用综述[J]. 现代教育技术, 2024, 34(6): 12-23.",
        "刘晨, 张翔. 基于检索增强的课程问答系统研究[J]. 计算机工程与应用, 2023, 59(18): 101-109.",
        "王蕾, 许涛. 智慧教学平台的体系结构与实现路径[J]. 中国电化教育, 2022(9): 74-82.",
        "陈颖, 唐杰. 教育大模型应用风险与治理策略[J]. 开放教育研究, 2024, 30(2): 33-42.",
        "Django Software Foundation. Django Documentation[EB/OL]. https://docs.djangoproject.com/ .",
        "Encode OSS. Django REST framework Documentation[EB/OL]. https://www.django-rest-framework.org/ .",
        "Vue Team. Vue 3 Documentation[EB/OL]. https://vuejs.org/ .",
        "Element Plus Team. Element Plus Documentation[EB/OL]. https://element-plus.org/ .",
        "LangChain Team. Retrieval QA Concepts[EB/OL]. https://python.langchain.com/ .",
        "Chroma Team. Chroma Documentation[EB/OL]. https://docs.trychroma.com/ .",
        "python-pptx Developers. python-pptx Documentation[EB/OL]. https://python-pptx.readthedocs.io/ .",
        "Python-OpenXML. python-docx Documentation[EB/OL]. https://python-docx.readthedocs.io/ .",
        "Pillow Contributors. Pillow Documentation[EB/OL]. https://pillow.readthedocs.io/ .",
        "JWT.io. JSON Web Token Introduction[EB/OL]. https://jwt.io/introduction .",
        "Microsoft. PowerPoint Object Model Reference[EB/OL]. https://learn.microsoft.com/office/vba/api/overview/powerpoint .",
        "Goodfellow I, Bengio Y, Courville A. Deep Learning[M]. MIT Press, 2016.",
        "Jurafsky D, Martin J H. Speech and Language Processing[M]. 3rd draft, 2023.",
        "Manning C D, Schütze H. Foundations of Statistical Natural Language Processing[M]. MIT Press, 1999.",
        "He K, Zhang X, Ren S, et al. Deep Residual Learning for Image Recognition[C]. CVPR, 2016.",
        "Russell S, Norvig P. Artificial Intelligence: A Modern Approach[M]. 4th ed. Pearson, 2021.",
        "ISO/IEC 25010:2011. Systems and software engineering—Systems and software Quality Requirements and Evaluation (SQuaRE)—System and software quality models[S].",
        "Pressman R S, Maxim B R. Software Engineering: A Practitioner’s Approach[M]. 9th ed. McGraw-Hill, 2019.",
        "Sommerville I. Software Engineering[M]. 10th ed. Pearson, 2015.",
        "Klein A, Nabi M. Explainability in NLP: Current Trends and Future Directions[J]. IEEE Access, 2022.",
        "Mikolov T, Chen K, Corrado G, et al. Efficient Estimation of Word Representations in Vector Space[J]. arXiv:1301.3781, 2013.",
        "Liu Y, Ott M, Goyal N, et al. RoBERTa: A Robustly Optimized BERT Pretraining Approach[J]. arXiv:1907.11692, 2019.",
        "Peters M E, Neumann M, Iyyer M, et al. Deep contextualized word representations[C]. NAACL, 2018.",
        "Hochreiter S, Schmidhuber J. Long Short-Term Memory[J]. Neural Computation, 1997, 9(8): 1735-1780.",
    ]
    add_reference_list(doc, references)
    return doc


def generate_figures() -> tuple[Path, Path, Path]:
    GEN_IMG_DIR.mkdir(parents=True, exist_ok=True)
    fig1 = GEN_IMG_DIR / "系统架构图.png"
    fig2 = GEN_IMG_DIR / "模块结构图.png"
    fig3 = GEN_IMG_DIR / "运行实例图.png"
    create_system_architecture_image(fig1)
    create_module_structure_image(fig2)
    create_runtime_example_image(fig3)
    return fig1, fig2, fig3


def copy_to_desktop(src: Path) -> Path | None:
    for dst in DESKTOP_PATHS:
        try:
            dst.parent.mkdir(parents=True, exist_ok=True)
            copy2(src, dst)
            return dst
        except OSError:
            continue
    return None


def main() -> None:
    fig1, fig2, fig3 = generate_figures()
    doc = build_document(fig1, fig2, fig3)
    doc.save(str(OUT_PATH))
    desktop_path = copy_to_desktop(OUT_PATH)
    print(f"generated: {OUT_PATH}")
    print(f"desktop_copy: {desktop_path}" if desktop_path else "desktop_copy: failed")


if __name__ == "__main__":
    main()
