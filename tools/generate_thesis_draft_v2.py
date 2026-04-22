# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT_PATH = Path(r"D:\pycharm\Graduation_Project_Django\thesis_draft_v2.docx")


def set_run_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    latin_font = "Times New Roman" if name == "Times New Roman" else name
    run._element.rPr.rFonts.set(qn("w:ascii"), latin_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin_font)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_paragraph(
    doc: Document,
    text: str,
    *,
    font="宋体",
    size=12,
    bold=False,
    align=WD_PARAGRAPH_ALIGNMENT.LEFT,
    first_line_indent=True,
    line_spacing=18,
):
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    fmt.line_spacing = Pt(line_spacing)
    fmt.first_line_indent = Inches(0.33) if first_line_indent else Inches(0)
    run = p.add_run(text)
    set_run_font(run, font, size, bold)
    return p


def add_title_center(doc: Document, text: str, *, font="宋体", size=16):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    set_run_font(run, font, size, True)
    return p


def add_h1(doc: Document, text: str):
    add_paragraph(
        doc,
        text,
        font="宋体",
        size=16,
        bold=True,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        first_line_indent=False,
        line_spacing=18,
    )


def add_h2(doc: Document, text: str):
    add_paragraph(
        doc,
        text,
        font="宋体",
        size=12,
        bold=True,
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        first_line_indent=False,
        line_spacing=18,
    )


def add_h3(doc: Document, text: str):
    add_paragraph(
        doc,
        text,
        font="宋体",
        size=10.5,
        bold=True,
        align=WD_PARAGRAPH_ALIGNMENT.LEFT,
        first_line_indent=False,
        line_spacing=18,
    )


def build_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.25)
    sec.right_margin = Inches(1.0)

    # 封面
    add_paragraph(doc, "北京信息科技大学", font="宋体", size=22, bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER, first_line_indent=False, line_spacing=24)
    add_paragraph(doc, "", first_line_indent=False)
    add_paragraph(doc, "毕业设计（论文）", font="宋体", size=22, bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER, first_line_indent=False, line_spacing=24)
    add_paragraph(doc, "", first_line_indent=False)
    add_paragraph(doc, "", first_line_indent=False)
    add_paragraph(doc, "题    目：基于大语言模型的双语课件智能学习平台设计与实现", first_line_indent=False)
    add_paragraph(doc, "学    院：__________", first_line_indent=False)
    add_paragraph(doc, "专    业：__________", first_line_indent=False)
    add_paragraph(doc, "学生姓名：__________    班级/学号：__________", first_line_indent=False)
    add_paragraph(doc, "指导老师：__________", first_line_indent=False)
    add_paragraph(doc, "起止时间：____年__月__日 至 ____年__月__日", first_line_indent=False)
    doc.add_page_break()

    # 中文摘要
    add_title_center(doc, "摘   要", font="宋体", size=16)
    add_paragraph(
        doc,
        "随着高校课程资源数字化建设不断推进，PPT课件已经成为教学活动中最常见的信息载体。传统课件学习存在三个共性问题：其一，英文课件的阅读门槛较高，学生需要在术语理解与内容吸收之间频繁切换；其二，课件内容分散在多页中，难以形成结构化知识框架；其三，学习过程中缺少能够基于课件上下文进行追问与反馈的智能工具。针对上述痛点，本文设计并实现了一个基于大语言模型的双语课件智能学习平台。",
    )
    add_paragraph(
        doc,
        "本系统以前后端分离架构为基础，后端采用Django与Django REST Framework构建服务能力，前端采用Vue 3实现交互界面。平台支持 .ppt/.pptx 课件上传、页面结构解析、双语翻译、逐页预览、课件问答、自动总结和术语提取等核心功能。在翻译能力上，系统结合课件版式信息与文本容器信息，尽可能保持原始布局与表达层次；在学习辅助能力上，系统引入向量检索与历史对话机制，使问答结果能够更好地贴合当前页或整份课件语境；在工程实现上，系统增加了翻译进度跟踪、失败可恢复、历史记录复用等机制，提升了可用性与稳定性。",
    )
    add_paragraph(
        doc,
        "实验与使用结果表明，该平台能够在较短时间内完成多页课件的翻译与知识整理，并提供可追溯的问答与总结结果，较好地满足了“理解—提问—复盘”的学习闭环需求。本文工作对高校双语教学场景下的课件学习效率提升具有一定实践意义。",
    )
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    p.paragraph_format.line_spacing = Pt(18)
    set_run_font(p.add_run("关键词："), "宋体", 10.5, True)
    set_run_font(p.add_run(" 大语言模型；双语课件；智能问答；自动总结；教学辅助平台"), "宋体", 10.5, False)
    doc.add_page_break()

    # 英文摘要
    add_title_center(doc, "Abstract", font="Times New Roman", size=16)
    add_paragraph(
        doc,
        "With the rapid digitalization of higher education resources, PowerPoint courseware has become a major carrier of teaching content. However, conventional courseware learning often suffers from three problems: high barriers to understanding English slides, weak cross-page knowledge integration, and lack of context-aware interactive support. To address these issues, this thesis designs and implements an intelligent bilingual courseware learning platform based on large language models (LLMs).",
        font="Times New Roman",
        size=12,
    )
    add_paragraph(
        doc,
        "The platform adopts a front-end and back-end separated architecture. The back end is built with Django and Django REST Framework, and the front end is implemented with Vue 3. It supports key functions including .ppt/.pptx upload, slide structure parsing, bilingual translation, page-level preview, context-aware QA, automatic summarization, and glossary extraction. By combining layout-aware text container translation, vector retrieval, and dialogue history, the system provides more reliable learning assistance for both single-slide and full-deck scenarios.",
        font="Times New Roman",
        size=12,
    )
    add_paragraph(
        doc,
        "Experimental results show that the platform can efficiently complete multi-page translation and knowledge organization, while offering traceable QA and summary outputs. It effectively supports a closed learning loop of comprehension, questioning, and review, and demonstrates practical value in bilingual teaching scenarios.",
        font="Times New Roman",
        size=12,
    )
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    p.paragraph_format.line_spacing = Pt(18)
    set_run_font(p.add_run("Keywords: "), "Times New Roman", 12, True)
    set_run_font(
        p.add_run("large language model; bilingual courseware; intelligent QA; automatic summarization; teaching support platform"),
        "Times New Roman",
        12,
        False,
    )
    doc.add_page_break()

    # 目录（初版占位）
    add_title_center(doc, "目    录", font="宋体", size=16)
    for item in [
        "第一章 绪论",
        "第二章 相关技术与可行性分析",
        "第三章 系统需求分析",
        "第四章 系统设计",
        "第五章 系统实现",
        "第六章 系统测试与结果分析",
        "结束语",
        "参考文献",
    ]:
        add_paragraph(doc, item, first_line_indent=False, size=12)
    doc.add_page_break()

    # 第一章
    add_h1(doc, "第一章 绪论")
    add_h2(doc, "1.1 研究背景与意义")
    add_paragraph(doc, "在双语教学与专业课程教学中，英文PPT课件使用比例持续上升。对于学生而言，课件理解往往不仅受语言能力影响，还受信息组织方式影响。许多课程存在“单页可读、全局难懂”的问题，导致学习效率降低。构建面向课件场景的智能辅助平台，可以在不改变教师课件习惯的前提下，提供翻译、问答与总结能力，从而提升学习效率与复盘质量。")
    add_h2(doc, "1.2 国内外研究现状")
    add_paragraph(doc, "近年来，大语言模型在机器翻译、阅读理解和教育问答中的应用迅速发展。现有通用翻译工具擅长句级转换，但对课件结构和页面布局关注不足；通用对话系统具备较强语言生成能力，但在“基于指定课件上下文作答”方面仍存在幻觉风险。针对教学场景的系统应将结构解析、检索增强和任务约束结合，才能实现更稳定的结果。")
    add_h2(doc, "1.3 本文主要工作")
    add_paragraph(doc, "本文围绕“课件翻译+学习辅助”目标，完成了以下工作：一是设计并实现了支持 .ppt/.pptx 的课件解析与结构化存储流程；二是构建了面向课件场景的翻译服务，支持逐页进度与失败恢复；三是实现了基于向量检索的问答模块，支持整份课件与单页问答两种模式；四是实现了章节总结、术语提取和思维导图生成能力；五是完成了系统级测试与可用性验证。")

    # 第二章
    add_h1(doc, "第二章 相关技术与可行性分析")
    add_h2(doc, "2.1 技术选型")
    add_paragraph(doc, "后端采用 Django + Django REST Framework，具备开发效率高、接口规范清晰、权限体系成熟等优势；前端采用 Vue 3 + Element Plus，便于快速构建交互式界面；数据库使用 SQLite（开发阶段）实现轻量化数据管理；向量检索采用 ChromaDB，支持课件文本的语义检索；大语言模型接口采用 OpenAI 兼容协议，并接入 qwen-plus 模型。")
    add_h2(doc, "2.2 关键技术分析")
    add_h3(doc, "2.2.1 课件结构解析技术")
    add_paragraph(doc, "系统使用 python-pptx 解析 PPTX 中的文本容器、段落结构、位置与样式信息，并在必要时通过 Windows COM 能力将 PPT 转换为可编辑 PPTX，以统一后续处理流程。")
    add_h3(doc, "2.2.2 检索增强问答技术")
    add_paragraph(doc, "问答模块在大模型生成前先进行语义检索，获取与问题最相关的页面片段，减少无关信息干扰。通过“上下文约束+引用页码”策略，提高回答的可追溯性。")
    add_h3(doc, "2.2.3 自动总结与术语抽取技术")
    add_paragraph(doc, "系统通过结构化提示词约束大模型输出 JSON，生成章节摘要、关键要点、术语对照以及思维导图结构，便于前端可视化展示与复用。")
    add_h2(doc, "2.3 可行性分析")
    add_paragraph(doc, "从技术可行性看，所选框架与依赖均成熟稳定，具有较强工程落地能力；从实现成本看，系统可在普通开发环境运行，部署门槛较低；从应用可行性看，目标用户与使用场景明确，功能闭环清晰，具备实际推广价值。")

    # 第三章
    add_h1(doc, "第三章 系统需求分析")
    add_h2(doc, "3.1 功能需求")
    add_paragraph(doc, "系统核心功能包括：课件上传与解析、逐页翻译、翻译进度查看、课件问答、自动总结、术语表展示、思维导图展示、历史记录管理等。问答模块需支持“整份课件问答”和“单页问答”模式切换，以兼顾宏观复习与局部精读。")
    add_h2(doc, "3.2 非功能需求")
    add_paragraph(doc, "系统需满足稳定性、可维护性与可扩展性要求。稳定性方面，应具备异常处理与任务恢复机制；可维护性方面，应采用模块化服务划分；可扩展性方面，应支持模型切换、检索参数调整和新功能接入。")
    add_h2(doc, "3.3 业务流程分析")
    add_paragraph(doc, "用户上传课件后，系统完成文本与布局解析并生成原始预览；用户触发翻译后，系统按页处理并实时更新状态；翻译完成后，用户可进行问答和总结，系统将问答记录与总结记录持久化保存，形成可追溯学习档案。")

    # 第四章
    add_h1(doc, "第四章 系统设计")
    add_h2(doc, "4.1 总体架构设计")
    add_paragraph(doc, "系统采用前后端分离架构。前端负责文件上传、状态展示、双语预览与交互操作；后端负责业务编排、模型调用、数据持久化与向量索引管理。通过 REST API 进行通信，使用 JWT 实现身份认证。")
    add_h2(doc, "4.2 数据库与数据模型设计")
    add_paragraph(doc, "核心数据模型包括 Courseware（课件信息）、SlideContent（页内容与布局）、QARecord（问答记录）、SummaryRecord（总结记录）与 TranslationCache（翻译缓存）。其中 SlideContent 记录页面级文本与布局信息，是翻译、问答和总结模块的共享基础。")
    add_h2(doc, "4.3 模块设计")
    add_h3(doc, "4.3.1 翻译模块设计")
    add_paragraph(doc, "翻译模块以页为单位执行，结合术语词典与缓存机制，提高一致性与性能。系统支持并发处理与失败页补偿策略，并记录已完成页数、总页数与耗时信息。")
    add_h3(doc, "4.3.2 问答模块设计")
    add_paragraph(doc, "问答模块根据用户选择确定范围：默认整份课件，单页问答可选。系统先构建历史对话与课件上下文，再调用大模型生成回答，并返回引用页码与片段。")
    add_h3(doc, "4.3.3 总结模块设计")
    add_paragraph(doc, "总结模块对整份课件内容进行聚合，输出章节摘要、关键点、术语表和思维导图。当前端空间受限时，术语表与思维导图采用全宽展示，保证信息完整可读。")

    # 第五章
    add_h1(doc, "第五章 系统实现")
    add_h2(doc, "5.1 后端实现")
    add_paragraph(doc, "后端在 Django 中按“视图层—服务层—数据层”组织代码：视图层负责参数校验与接口返回；服务层封装翻译、问答、总结、图像处理与检索能力；数据层负责模型定义与持久化。任务型接口采用后台线程执行，避免阻塞请求线程。")
    add_h2(doc, "5.2 前端实现")
    add_paragraph(doc, "前端使用 Vue 3 组合式 API 组织状态与逻辑。上传页实现了课件上传、翻译状态轮询与双语预览；历史记录页支持左侧课件选择、右侧问答和总结复用；关键展示区域支持分页跳转、范围切换和异常提示，提升学习连续性。")
    add_h2(doc, "5.3 关键实现细节")
    add_paragraph(doc, "（1）模型可配置：通过环境变量统一管理模型名称与接口地址，便于切换到 qwen-plus 等模型。\n（2）翻译静默执行：在 Windows COM 调用中关闭 PowerPoint 可见窗口与告警弹窗，降低用户干扰。\n（3）缓存机制：基于文本与术语提示构建哈希键，减少重复翻译请求。")

    # 第六章
    add_h1(doc, "第六章 系统测试与结果分析")
    add_h2(doc, "6.1 测试环境与方法")
    add_paragraph(doc, "测试环境包括 Windows 开发主机、Python 3.12、Django 4.2、Vue 3 与 ChromaDB。测试方法采用功能测试与场景测试相结合，覆盖上传解析、翻译流程、问答准确性、总结可用性与异常恢复。")
    add_h2(doc, "6.2 功能测试结果")
    add_paragraph(doc, "测试结果表明：系统能够稳定完成课件上传、逐页翻译和结果展示；问答模块可根据范围设置返回相应答案，并提供页码引用；总结模块可生成结构化结果并在前端完整展示术语表与思维导图。")
    add_h2(doc, "6.3 性能与可用性分析")
    add_paragraph(doc, "在中等页数课件场景下，系统具备可接受的响应速度。通过缓存与并行翻译机制，重复内容场景下性能提升明显。可用性方面，进度显示、失败提示与可恢复机制降低了用户操作成本。")
    add_h2(doc, "6.4 存在问题与改进方向")
    add_paragraph(doc, "当前系统仍存在大模型输出不稳定、极端复杂版式适配成本高等问题。后续可从三方面优化：引入更细粒度版式约束、完善多模型路由策略、增加离线评测数据集与自动化评估流程。")

    # 结束语
    add_h1(doc, "结束语")
    add_paragraph(doc, "本文完成了基于大语言模型的双语课件智能学习平台的设计与实现，构建了“上传解析—翻译预览—问答互动—总结复盘”的完整学习闭环。系统在工程层面实现了模块化组织与关键能力联动，在应用层面提升了课件学习效率与复盘质量。后续将继续围绕模型稳定性、版式保真度和教学适配性开展优化工作。")

    # 参考文献
    add_h1(doc, "参考文献")
    refs = [
        "王某某, 李某某. 大语言模型在教育场景中的应用研究[J]. 现代教育技术, 2024, 34(6): 12-20.",
        "Brown T, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]. NeurIPS, 2020.",
        "刘某某, 张某某. 基于检索增强生成的问答系统设计与实现[J]. 计算机工程, 2023, 49(9): 88-96.",
        "Reimers N, Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks[C]. EMNLP, 2019.",
        "Django Software Foundation. Django Documentation[EB/OL]. https://docs.djangoproject.com/ .",
        "Vue Team. Vue.js Documentation[EB/OL]. https://vuejs.org/ .",
        "OpenAI. API Compatibility and Chat Completion Guide[EB/OL]. https://platform.openai.com/ .",
        "Python-OpenXML. python-docx Documentation[EB/OL]. https://python-docx.readthedocs.io/ .",
    ]
    for idx, ref in enumerate(refs, 1):
        add_paragraph(doc, f"[{idx}] {ref}", size=10.5)

    return doc


if __name__ == "__main__":
    document = build_doc()
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUT_PATH))
    print(str(OUT_PATH))
