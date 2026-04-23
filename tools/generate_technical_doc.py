# -*- coding: utf-8 -*-
from __future__ import annotations

import ast
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "技术.docx"

EXCLUDE_DIR_NAMES = {
    ".git",
    ".venv",
    "node_modules",
    "media",
    "dist",
    "__pycache__",
    "chroma_db",
    ".idea",
    ".cursor",
    ".pytest_cache",
}

EXCLUDE_FILE_SUFFIXES = {
    ".docx",
    ".pdf",
    ".ppt",
    ".pptx",
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".svg",
    ".ico",
    ".sqlite3",
    ".db",
    ".log",
    ".bin",
    ".pickle",
    ".pyc",
}

TEXT_EXTENSIONS = {
    ".py",
    ".pyw",
    ".vue",
    ".ts",
    ".js",
    ".css",
    ".html",
    ".md",
    ".txt",
    ".bat",
    ".json",
    ".d.ts",
}

TEXT_FILENAMES = {
    ".env",
    ".env.example",
    "requirements.txt",
    "package.json",
    "package-lock.json",
    "tsconfig.json",
    "vite.config.ts",
    "README.md",
}

FILE_NOTE_OVERRIDES: dict[str, dict[str, str]] = {
    "backend/learning/services/translation_service.py": {
        "role": (
            "该文件是整套课件翻译流程的总控中枢，负责把“原始课件文本 / OCR 文本 / 备注文本”"
            "组织成可翻译的结构化输入，再把模型返回结果重新映射回页面布局。它不只是调用一次大模型，"
            "而是同时承担术语提示拼装、缓存命中、分块翻译、失败重试、PDF 重复短语去重兼容、"
            "翻译结果持久化以及译后预览触发等职责。"
        ),
        "tech": (
            "Django ORM、OpenAI 兼容接口、JSON 结构化输出约束、SHA256 缓存键、"
            "ThreadPoolExecutor 并发调度、版式感知翻译、OCR 文本补充。"
        ),
        "hard": (
            "难点在于同时满足翻译质量、术语一致性、吞吐效率和版式可回填性。"
            "如果只做纯文本翻译，结果容易脱离页面布局；如果过度依赖结构化输出，又会因为模型漏项导致失败。"
            "因此这里采用“缓存优先 -> 结构化翻译 -> 缺失重试 -> 单条兜底 -> 持久化”的多层策略。"
        ),
    },
    "backend/learning/services/ppt_parser_service.py": {
        "role": (
            "该文件负责把 PPT/PDF 课件解析成统一的内部数据结构。核心目标不是简单抽文本，"
            "而是保留页面尺寸、块坐标、段落顺序、字体大小、标题候选、表格单元格以及图片 OCR 占位信息，"
            "为后续翻译回填、预览渲染和问答检索提供稳定输入。"
        ),
        "tech": (
            "python-pptx、PyMuPDF、Windows COM、分组形状递归遍历、容器排序、"
            "PDF 文本块抽取、短语级去重与标题识别。"
        ),
        "hard": (
            "难点在于 PPT 与 PDF 的结构差异极大：PPT 是形状树，PDF 更接近文本块坐标流。"
            "解析层必须把异构输入统一成同一种布局模型，同时过滤页眉页脚、水印、重复短语等噪声，"
            "否则后续翻译和导出都会出现错位或重复。"
        ),
    },
    "backend/learning/services/image_processing_service.py": {
        "role": (
            "该文件负责译后可视化落地，把翻译结果真正写回 PPT 或 PDF 预览图中。"
            "它既要处理原生 PPT 文本框回填，也要处理 PDF 页面图片化后的文字覆盖，"
            "并尽量保持原有版式、字号层级和可读性。"
        ),
        "tech": (
            "python-pptx、Pillow、字体候选回退、字号二分拟合、背景采样、对比度计算、"
            "PDF 页面重绘与导出。"
        ),
        "hard": (
            "难点在于译文长度通常比原文更不稳定，尤其在固定文本框内很容易溢出。"
            "因此这里专门把“字体选择、字号拟合、颜色对比、长页压缩、PDF 图片重绘”拆成独立函数，"
            "以便逐层控制可读性。"
        ),
    },
    "backend/learning/services/qa_service.py": {
        "role": (
            "该文件负责课件问答的上下文组织与提示词拼装。它将向量检索结果、当前页内容、"
            "历史对话和作用范围（单页/整份课件）组合成可追溯的问答上下文，并返回带引用的答案。"
        ),
        "tech": "RAG、向量检索、上下文裁剪、对话历史拼接、引用回传。",
        "hard": (
            "难点不在于单纯生成答案，而在于限制模型只围绕课件内容回答，并保证引用页码可追踪。"
            "如果上下文组织不稳，问答很容易变成泛化回答。"
        ),
    },
    "backend/learning/services/vector_index_service.py": {
        "role": (
            "该文件负责把课件内容写入向量库并提供检索接口，是问答与知识定位能力的基础设施层。"
            "它需要为每条内容生成稳定文档 ID、维护课件级过滤条件，并在重建索引时清理旧数据。"
        ),
        "tech": "Chroma PersistentClient、文档 ID 规则、where 过滤、相似度查询。",
        "hard": (
            "难点在于索引一致性和重建成本。若文档 ID 设计不稳定，就会在重复翻译或重复建库时产生脏数据。"
        ),
    },
    "backend/learning/services/summary_service.py": {
        "role": (
            "该文件负责把整份课件凝练为章节点摘要、知识要点、术语对照表、思维导图和学习建议。"
            "它既负责调用模型，也负责在模型输出不稳定时做 JSON 提取、结构修正和本地兜底。"
        ),
        "tech": "结构化摘要提示词、JSON 解析、思维导图归一化、本地降级总结。",
        "hard": (
            "难点在于模型输出格式不总是稳定，尤其是思维导图层级和建议列表格式容易漂移。"
            "所以这里必须保证“即使模型不完美，前端也能拿到完整结构”。"
        ),
    },
    "backend/learning/services/llm_client.py": {
        "role": (
            "该文件是大模型调用的统一封装层，向上屏蔽鉴权、接口地址、消息序列化、重试、"
            "环境变量刷新和异常分类等细节。业务层只关心传入消息与取回结果，不直接面对底层 HTTP 波动。"
        ),
        "tech": "requests、指数退避、dotenv 刷新、异常分层、消息序列化。",
        "hard": (
            "难点在于外部模型服务具有不确定性，既可能超时，也可能鉴权失效，还可能返回非预期结构。"
            "统一客户端层是保证全局稳定性的关键。"
        ),
    },
    "backend/learning/views.py": {
        "role": (
            "该文件是后端接口编排层，负责把上传、翻译、导出、问答、总结、历史记录等完整业务流程"
            "暴露为 API，并处理权限校验、状态流转、后台线程任务与响应封装。"
        ),
        "tech": "DRF APIView、事务控制、后台线程、FileResponse、状态聚合。",
        "hard": (
            "难点在于长耗时翻译任务不能阻塞请求线程，同时前端又必须持续拿到可观察的状态。"
            "因此这里把任务启动、状态轮询、后处理和导出分成多层接口。"
        ),
    },
    "backend/learning/models.py": {
        "role": (
            "该文件定义整套系统的核心数据模型，包括课件、页内容、问答记录、总结记录、术语词典和翻译缓存。"
            "它承担数据组织、关联关系和持久化边界的职责。"
        ),
        "tech": "Django ORM、JSONField、索引、唯一约束。",
        "hard": (
            "难点在于既要支持页面级细粒度查询，又要支持较大的布局 JSON 与缓存数据持久化。"
        ),
    },
    "backend/learning/serializers.py": {
        "role": (
            "该文件负责 API 输入输出校验，是请求进入业务层前的第一道边界控制。"
            "它统一约束上传文件、问答参数、记录结构和字段裁剪规则。"
        ),
        "tech": "DRF Serializer、ModelSerializer、字段验证、结构校验。",
        "hard": (
            "难点在于前端交互场景多，参数结构复杂，必须在序列化层尽早拦截脏数据，"
            "避免异常一路传到服务层。"
        ),
    },
    "frontend/src/views/UploadTranslateView.vue": {
        "role": (
            "这是前端主工作台页面，承载上传、翻译轮询、双语对照、备注翻译、问答、摘要生成、"
            "导出下载等完整主流程，是用户使用频率最高、状态最复杂的页面。"
        ),
        "tech": "Vue3 Composition API、Element Plus、Axios、Markdown 渲染、轮询状态同步。",
        "hard": (
            "难点在于一个页面内同时维护上传状态、翻译状态、页切换状态、问答状态和摘要状态，"
            "任何一个链路处理不好都会影响整体体验。"
        ),
    },
    "frontend/src/views/RecordsView.vue": {
        "role": (
            "该页面负责历史记录回看与复用，集中展示历史问答、总结记录、术语与预览内容，"
            "帮助用户从“当前工作模式”切换到“历史复盘模式”。"
        ),
        "tech": "Vue3、路由联动、Markdown 渲染、历史记录聚合。",
        "hard": "难点在于保持历史上下文与当前课件选择同步，同时避免页面切换后状态混乱。",
    },
    "frontend/src/views/DashboardLayout.vue": {
        "role": (
            "该文件提供前端整体布局壳层，统一负责顶部导航、课件选择器、子路由切换和全局事件广播。"
        ),
        "tech": "Vue Router、keep-alive、本地存储、全局事件。",
        "hard": "难点在于跨页面共享当前课件状态，并保证工作台与历史页的切换连续性。",
    },
    "frontend/src/router/index.ts": {
        "role": "该文件定义前端路由表和登录守卫，确保用户访问路径与登录状态一致。",
        "tech": "vue-router、前置守卫、路由重定向。",
        "hard": "难点在于未登录拦截和已登录回流逻辑必须稳定，否则会出现死循环跳转。",
    },
    "frontend/src/api/client.ts": {
        "role": (
            "该文件封装前端统一 HTTP 客户端，负责自动携带 JWT、拦截 401 并清理本地登录状态。"
        ),
        "tech": "Axios 实例、请求拦截器、响应拦截器。",
        "hard": "难点在于鉴权过期时需要立即恢复到明确的登录态，不能让页面停留在半失效状态。",
    },
    "launch_platform.pyw": {
        "role": (
            "该文件是本地桌面启动器，负责一键拉起前后端开发环境并实时显示日志，"
            "降低手动打开多个终端的成本。"
        ),
        "tech": "Tkinter、subprocess、线程日志读取。",
        "hard": "难点在于多进程生命周期管理和日志流式输出，尤其要保证关闭时能干净退出。",
    },
}

FUNCTION_NOTE_OVERRIDES: dict[str, dict[str, dict[str, str]]] = {
    "backend/learning/services/translation_service.py": {
        "TranslationService.translate_courseware": {
            "purpose": "作为整份课件翻译入口，负责顺序/并发翻译、异常收集、状态更新和最终结果落库。",
            "reason": "把整份课件的流程总控集中在这里，可以保证状态迁移、错误汇总和性能策略只维护一处。"
        },
        "TranslationService._build_translated_layout": {
            "purpose": "把源布局、文本容器、OCR 结果和翻译结果重新拼装成可回填的 translated_layout。",
            "reason": "翻译不是返回一段纯文本就结束，必须重新对齐到容器坐标，否则预览图与导出文件无法复原版式。"
        },
        "TranslationService._translate_containers_structured": {
            "purpose": "以结构化 JSON 的形式批量翻译多个文本容器，并尽量保持容器边界不丢失。",
            "reason": "相较逐条翻译，结构化批量翻译更高效，也更容易让模型理解多个容器之间的边界。"
        },
        "TranslationService._translate_container_chunk": {
            "purpose": "负责单个分块的完整翻译闭环，包括结构化翻译、重试和单容器兜底。",
            "reason": "把分块闭环单独抽出来后，整体翻译器只需要关注调度，不必混杂分块内部容错细节。"
        },
        "TranslationService._translate_image_containers": {
            "purpose": "处理图片型容器的 OCR 与翻译，将纯文本容器之外的信息补齐。",
            "reason": "课件中很多关键信息会直接嵌在图片里，如果不单独处理，用户看到的内容会不完整。"
        },
        "TranslationService._split_to_match_paragraphs": {
            "purpose": "把模型返回文本重新切回与原容器段落数量尽量一致的结构。",
            "reason": "段落数直接影响回填后的换行和版式，如果不做这一步，译文极易出现整块挤压或错位。"
        },
        "TranslationService._build_cache_key": {
            "purpose": "根据模型、术语、翻译类型和原文共同构造唯一缓存键。",
            "reason": "单纯用原文做缓存键会误命中不同术语上下文或模型版本下的旧结果，因此必须把影响翻译结果的因素一起纳入。"
        },
    },
    "backend/learning/services/ppt_parser_service.py": {
        "PPTParserService.parse_pdf": {
            "purpose": "把 PDF 页面拆成统一布局模型，并抽取文本块、图片块和标题候选。",
            "reason": "PDF 无法像 PPT 一样直接编辑，所以必须先把它转成统一内部布局，后续翻译和导出才能复用同一套流程。"
        },
        "PPTParserService.dedupe_pdf_repeated_short_phrases": {
            "purpose": "识别 PDF 中高频重复的短语型文本块，并只保留一次进入翻译流程。",
            "reason": "PDF 水印常以短语形式重复出现在多个位置，若不提前去重，翻译结果会充满重复内容，严重影响可读性。"
        },
        "PPTParserService._finalize_layout": {
            "purpose": "将容器排序、标题确定、block 重建和 page 元信息整合为最终 layout。",
            "reason": "解析阶段输出统一结构后，后续翻译、预览和向量索引都可以基于同一数据模型，降低链路耦合。"
        },
        "PPTParserService.parse_pptx": {
            "purpose": "解析 PPT/PPTX 并提取文本容器、表格、图片占位和备注信息。",
            "reason": "PPT 本身可编辑，解析时保留更多结构信息可以显著提高译后回填效果。"
        },
    },
    "backend/learning/services/image_processing_service.py": {
        "ImageProcessingService._set_text_frame_content": {
            "purpose": "把译文写回 PowerPoint 文本框，并自动调整字号、颜色、段落和边距。",
            "reason": "译文长度不可控，如果简单覆盖原文本，极容易溢出；这里集中处理排版适配能提高稳定性。"
        },
        "ImageProcessingService.export_translated_pdf": {
            "purpose": "将翻译后的 PDF 页面重新渲染成图片并组装成新的 PDF 文件。",
            "reason": "PPT 内部通常保留了真正可编辑的文本框，所以译文可以直接写回原文本对象；而 PDF 更接近已经排版完成的定版页面，文字对象不容易稳定修改，因此这里选择先按原坐标重画整页内容，再重新生成新的 PDF 文件，这样更稳。"
        },
        "ImageProcessingService._draw_paragraphs_on_image": {
            "purpose": "在 PDF 页面图片或预览图上绘制译文段落，并保持基础可读性。",
            "reason": "把图片绘制逻辑独立出来后，预览生成与正式 PDF 导出可以共用同一套版式策略。"
        },
        "ImageProcessingService._best_fit_font_size": {
            "purpose": "通过二分试探方式寻找当前文本框可容纳的最佳字号。",
            "reason": "逐级递减效率低且不稳定，二分拟合在保证结果可用的同时计算更快。"
        },
    },
    "backend/learning/services/qa_service.py": {
        "QAService.ask": {
            "purpose": "作为问答统一入口，负责调用检索、组织上下文、拼接历史对话并得到答案。",
            "reason": "把问答流程收敛为一个入口，便于统一控制作用范围、上下文模板和引用输出格式。"
        },
    },
    "backend/learning/services/summary_service.py": {
        "SummaryService.generate": {
            "purpose": "生成章节总结、要点、术语、思维导图等完整总结结构。",
            "reason": "总结结果是多字段组合输出，集中在一个函数里便于统一兜底策略和结构修正。"
        },
    },
    "backend/learning/services/llm_client.py": {
        "OpenAICompatibleClient.chat": {
            "purpose": "统一执行一次模型对话请求，并处理鉴权、重试、错误转换和内容反序列化。",
            "reason": "业务层不应该重复实现 HTTP 调用和异常处理，把调用入口统一起来能显著降低维护成本。"
        },
    },
    "backend/learning/views.py": {
        "TranslateCoursewareView.post": {
            "purpose": "启动课件翻译任务并立即返回可轮询状态，而不是阻塞等待整个翻译完成。",
            "reason": "翻译是长耗时操作，若不异步启动，前端会超时，用户也无法实时感知进度。"
        },
        "ExportTranslatedPPTView.get": {
            "purpose": "根据源文件类型导出翻译后的 PPTX 或 PDF，并以附件形式返回给前端。",
            "reason": "把导出入口统一在一个接口上，可以让前端只关心下载动作，不必分别判断文档类型。"
        },
    },
    "frontend/src/views/UploadTranslateView.vue": {
        "translateCurrent": {
            "purpose": "启动翻译并通过轮询持续刷新状态、页列表和错误提示。",
            "reason": "翻译是异步长任务，这里集中处理轮询和状态迁移，能避免逻辑散落在多个按钮事件中。"
        },
        "downloadTranslatedPpt": {
            "purpose": "下载翻译后的课件文件，并根据返回文件名修正下载类型与本地文件名。",
            "reason": "导出文件既可能是 PPTX 也可能是 PDF，前端必须根据响应动态兜底，避免浏览器错误识别类型。"
        },
        "refreshCoursewareStatus": {
            "purpose": "轮询后端获取当前课件的状态、进度、错误信息和更新时间。",
            "reason": "将状态刷新抽成单独函数，可以让首次加载、轮询和手动刷新复用同一条逻辑链。"
        },
        "uploadFile": {
            "purpose": "负责课件上传、重置工作区状态并加载新课件内容。",
            "reason": "上传成功后需要同步重置翻译、问答、摘要和预览等多个区域，必须集中处理。"
        },
    },
    "frontend/src/views/RecordsView.vue": {
        "loadAll": {
            "purpose": "并行加载当前课件的页数据、问答记录和总结记录。",
            "reason": "历史页依赖多类数据源，集中加载更容易控制加载顺序和错误反馈。"
        },
        "submitQuestion": {
            "purpose": "在历史页上下文中继续发起问答，并把答案写回消息列表。",
            "reason": "复盘场景下仍然需要延续式提问，因此这里保留与工作台相近的问答组织能力。"
        },
    },
    "launch_platform.pyw": {
        "LauncherApp.start_all": {
            "purpose": "一键启动前后端服务，降低本地运行门槛。",
            "reason": "开发和演示时最常见的操作是同时启动两个服务，因此提供合并入口可以减少误操作。"
        },
        "LauncherApp.stop_all": {
            "purpose": "统一停止前后端子进程并清理运行状态。",
            "reason": "集中关闭比依赖用户手动逐个关终端更可靠，也能减少僵尸进程。"
        },
    },
}


@dataclass
class FunctionDetail:
    symbol: str
    signature: str
    purpose: str
    reason: str


def is_excluded(path: Path) -> bool:
    return any(part in EXCLUDE_DIR_NAMES for part in path.parts)


def is_text_file(path: Path) -> bool:
    ext = path.suffix.lower()
    return ext in TEXT_EXTENSIONS or path.name in TEXT_FILENAMES or path.name.endswith(".d.ts")


def should_skip_file(path: Path) -> bool:
    return path.suffix.lower() in EXCLUDE_FILE_SUFFIXES


def safe_read_text(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "gbk", "cp936", "latin1"):
        try:
            return path.read_text(encoding=enc)
        except Exception:
            continue
    return ""


def format_python_signature(node: ast.FunctionDef) -> str:
    args: list[str] = []
    positional = list(node.args.posonlyargs) + list(node.args.args)
    defaults = [None] * (len(positional) - len(node.args.defaults)) + list(node.args.defaults)
    for arg, default in zip(positional, defaults):
        if arg.arg in {"self", "cls"}:
            continue
        if default is None:
            args.append(arg.arg)
        else:
            args.append(f"{arg.arg}=...")
    if node.args.vararg:
        args.append(f"*{node.args.vararg.arg}")
    if node.args.kwonlyargs:
        if not node.args.vararg:
            args.append("*")
        for arg, default in zip(node.args.kwonlyargs, node.args.kw_defaults):
            if default is None:
                args.append(arg.arg)
            else:
                args.append(f"{arg.arg}=...")
    if node.args.kwarg:
        args.append(f"**{node.args.kwarg.arg}")
    return f"({', '.join(args)})"


def extract_python_symbols(content: str) -> tuple[list[str], list[str], list[str]]:
    classes: list[str] = []
    funcs: list[str] = []
    imports: list[str] = []
    try:
        tree = ast.parse(content)
    except Exception:
        return classes, funcs, imports

    for node in tree.body:
        if isinstance(node, ast.ClassDef):
            classes.append(node.name)
        elif isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
            funcs.append(node.name)
        elif isinstance(node, ast.Import):
            imports.extend(alias.name.split(".")[0] for alias in node.names)
        elif isinstance(node, ast.ImportFrom) and node.module:
            imports.append(node.module.split(".")[0])
    return classes, funcs, sorted(set(imports))


def extract_ts_like_symbols(content: str) -> list[str]:
    result: list[str] = []
    patterns = [
        re.compile(r"^(?:export\s+)?function\s+([A-Za-z_][A-Za-z0-9_]*)\s*\("),
        re.compile(r"^(?:export\s+)?const\s+([A-Za-z_][A-Za-z0-9_]*)\s*=\s*(?:async\s*)?\([^)]*\)\s*=>"),
    ]
    for line in content.splitlines():
        raw = line.rstrip()
        if not raw or raw.startswith((" ", "\t")):
            continue
        for pattern in patterns:
            matched = pattern.search(raw)
            if matched:
                result.append(matched.group(1))
                break
    return result[:30]


def split_tokens(name: str) -> list[str]:
    normalized = re.sub(r"([a-z0-9])([A-Z])", r"\1_\2", name).strip("_").lower()
    return [token for token in normalized.split("_") if token]


def infer_function_purpose(relative_path: str, symbol_name: str) -> str:
    rp = relative_path.replace("\\", "/")
    base = symbol_name.split(".")[-1]
    tokens = split_tokens(base)

    if "translation_service.py" in rp:
        if "cache" in tokens and "key" in tokens:
            return "构造翻译缓存键，确保不同术语、模型和文本上下文被正确区分。"
        if "cache" in tokens and "lookup" in tokens:
            return "读取翻译缓存，避免重复调用模型。"
        if "cache" in tokens and "store" in tokens:
            return "写入翻译缓存，保存当前翻译结果供后续复用。"
        if "translate" in tokens and "image" in tokens:
            return "处理图片容器的 OCR 与翻译，将图片内文字纳入课件翻译范围。"
        if "translate" in tokens and "container" in tokens:
            return "围绕文本容器执行结构化翻译，并保持容器边界。"
        if "translate" in tokens:
            return "执行指定粒度的翻译任务，并返回可继续处理的结果。"
        if "layout" in tokens:
            return "将源布局与译文重新组装，形成可回填和可渲染的译后布局。"
        if "chunk" in tokens:
            return "按照上下文长度限制对容器进行分块，控制单次模型请求规模。"
        if "split" in tokens:
            return "按原段落结构切分译文，尽量保持原有阅读节奏。"
        if "preview" in tokens:
            return "生成译后预览图，方便前端即时查看效果。"

    if "ppt_parser_service.py" in rp:
        if "parse" in tokens:
            return "解析指定格式课件，并抽取统一布局结构。"
        if "extract" in tokens:
            return "从页面元素中提取局部结构信息，如文本、表格、图片或字体。"
        if "sort" in tokens:
            return "根据坐标和优先级排序，恢复接近人工阅读的内容顺序。"
        if "title" in tokens:
            return "识别或标记页面标题，提升后续展示与索引的语义性。"
        if "dedupe" in tokens or "repeated" in tokens:
            return "过滤 PDF 中重复出现的短语型文本块，减少水印噪声。"
        if "container" in tokens:
            return "生成统一容器结构，为后续翻译和渲染提供稳定输入。"

    if "image_processing_service.py" in rp:
        if "font" in tokens and "size" in tokens:
            return "计算当前内容可容纳的最佳字号，平衡完整显示与清晰度。"
        if "font" in tokens:
            return "选择或应用合适字体，保证多语言文字能够正常显示。"
        if "contrast" in tokens or "luminance" in tokens:
            return "计算文字与背景的对比关系，提升译后可读性。"
        if "background" in tokens or "sample" in tokens:
            return "估计原页面背景颜色，为译文覆盖提供视觉基础。"
        if "draw" in tokens or "render" in tokens:
            return "将译文绘制到图片或页面上，生成可见的译后结果。"
        if "export" in tokens:
            return "导出译后的 PPTX 或 PDF 文件，形成最终可下载成果。"
        if "process" in tokens:
            return "驱动整批页面的译后处理与预览输出。"
        if "apply" in tokens:
            return "把计算结果真正应用到具体页面元素上。"

    if "qa_service.py" in rp:
        if "context" in tokens:
            return "构建问答上下文，控制模型回答范围。"
        if "history" in tokens:
            return "整理历史问答内容，让模型理解对话连续性。"
        if "payload" in tokens:
            return "提取当前页核心内容，形成检索和问答输入。"
        if "ask" in tokens:
            return "发起一次问答请求并返回答案与引用。"

    if "summary_service.py" in rp:
        if "fallback" in tokens:
            return "在模型输出失败时生成可用的本地兜底摘要。"
        if "normalize" in tokens:
            return "修正模型输出结构，使前端可以稳定消费。"
        if "suggestion" in tokens:
            return "基于总结内容生成学习建议，补充面向学生的解释层。"
        if "generate" in tokens:
            return "组织整份课件的总结生成流程，并返回结构化结果。"

    if "vector_index_service.py" in rp:
        if "doc" in tokens and "id" in tokens:
            return "构造稳定的文档 ID，保证索引可重复构建。"
        if "where" in tokens:
            return "拼装向量检索过滤条件，限定查询范围。"
        if "rebuild" in tokens:
            return "重建课件向量索引，确保检索结果与最新内容一致。"
        if "query" in tokens:
            return "执行向量检索并返回最相关内容。"

    if "llm_client.py" in rp:
        if "retryable" in tokens:
            return "判断当前错误状态是否值得自动重试。"
        if "serialize" in tokens:
            return "将复杂消息内容整理为模型接口可接受的结构。"
        if "reload" in tokens or "dotenv" in tokens:
            return "重新加载环境配置，降低本地开发时的重启成本。"
        if "chat" in tokens:
            return "向模型服务发起对话请求并统一处理返回结果。"

    if rp.endswith("views.py"):
        if "progress" in tokens or "duration" in tokens:
            return "聚合课件状态指标，给前端提供可直接展示的数据。"
        if "post" in tokens or "get" in tokens:
            return "处理对应 API 请求，并连接业务服务层。"
        if "run" in tokens:
            return "在后台继续执行翻译后的后处理任务，避免阻塞主流程。"

    if rp.endswith(".vue") or rp.endswith(".ts") or rp.endswith(".js"):
        if "refresh" in tokens:
            return "刷新页面关键状态，使前端界面与后端进度保持同步。"
        if "load" in tokens or "fetch" in tokens or "hydrate" in tokens:
            return "加载页面所需数据，并在界面中完成状态落位。"
        if "upload" in tokens:
            return "处理上传动作及其后续状态初始化。"
        if "translate" in tokens:
            return "发起翻译相关交互，并联动页面状态变化。"
        if "download" in tokens:
            return "处理导出下载逻辑，保证文件在浏览器侧正确落地。"
        if "submit" in tokens:
            return "提交用户输入并对界面结果做更新。"
        if "handle" in tokens:
            return "响应用户事件或外部事件，协调局部状态变更。"
        if "jump" in tokens or "prev" in tokens or "next" in tokens:
            return "控制页面预览或记录浏览的导航行为。"

    if tokens[:1] == ["init"]:
        return "负责对象初始化，装配运行期依赖和默认参数。"
    if "normalize" in tokens:
        return "将输入数据转换为统一格式，降低后续处理分支复杂度。"
    if "build" in tokens:
        return "组装中间数据结构，供主流程复用。"
    if "extract" in tokens:
        return "提取局部信息，避免主流程函数承担过多解析细节。"
    if "resolve" in tokens:
        return "把外部表示映射为内部可操作对象。"
    if "parse" in tokens:
        return "解析输入内容并生成结构化结果。"
    if "mark" in tokens:
        return "标记当前对象或任务状态，便于后续流程判断。"
    if "save" in tokens or "store" in tokens:
        return "保存中间结果或最终结果，保证流程可追溯。"
    if "run" in tokens or "process" in tokens:
        return "承担一段完整的流程编排逻辑。"
    return "承担当前模块中的一个清晰子步骤，减少主流程函数的复杂度。"


def infer_function_reason(relative_path: str, symbol_name: str) -> str:
    rp = relative_path.replace("\\", "/")
    base = symbol_name.split(".")[-1]
    tokens = split_tokens(base)

    if "cache" in tokens:
        return "将缓存逻辑单独拆分，可以把性能优化和业务翻译逻辑解耦，后续替换缓存策略时影响更小。"
    if "normalize" in tokens:
        return "单独设置归一化函数可以集中处理脏数据和边界输入，避免这些细节散落在主流程中。"
    if "build" in tokens:
        return "通过独立构造函数统一数据拼装口径，能够减少重复代码并保证多个调用点结果一致。"
    if "extract" in tokens:
        return "把抽取逻辑细分后，解析异常更容易定位，也方便未来按格式继续扩展。"
    if "split" in tokens or "chunk" in tokens:
        return "分块或切分策略单独存在，便于围绕模型上下文长度和页面版式进行持续优化。"
    if "translate" in tokens and "courseware" not in tokens:
        return "将不同粒度的翻译动作拆开，可以在质量、性能和容错之间分别调优，而不是所有情况都走同一条路径。"
    if "translate" in tokens and "courseware" in tokens:
        return "总控函数单独存在，可以统一管理整份课件的状态变化、并发策略和错误汇总。"
    if "parse" in tokens:
        return "解析函数与业务编排分开后，格式差异只会影响解析层，不会向翻译和问答层扩散。"
    if "export" in tokens:
        return "导出逻辑独立可以把‘页面渲染’与‘文件封装’分层，后续新增格式时改动范围更可控。"
    if "query" in tokens or "ask" in tokens:
        return "查询入口单独定义后，上下文组织、引用策略和模型提示可以保持统一。"
    if "generate" in tokens:
        return "生成类函数集中控制输出格式，有利于在模型不稳定时统一加入兜底和修正策略。"
    if rp.endswith(".vue") or rp.endswith(".ts") or rp.endswith(".js"):
        return "前端交互函数拆分得足够细，可以把数据请求、状态更新和界面反馈分离，降低单个页面函数的复杂度。"
    if rp.endswith("views.py"):
        return "接口层函数按动作拆开后，更容易看清每个 API 的职责边界，也便于后续做权限和日志扩展。"
    if "font" in tokens or "contrast" in tokens or "background" in tokens:
        return "把排版和视觉相关决策拆成多个小函数，可以针对不同文档类型逐步调优，不必每次修改整条导出链路。"
    if "mark" in tokens:
        return "状态写入函数独立存在，可以保证数据库更新口径统一，减少不同分支写出不一致状态的风险。"
    return "将这个子步骤拆开后，主流程会更短、更可测，也更方便针对单一问题做局部优化。"


def build_function_detail(relative_path: str, symbol_name: str, signature: str) -> FunctionDetail:
    override = FUNCTION_NOTE_OVERRIDES.get(relative_path, {}).get(symbol_name)
    if override:
        return FunctionDetail(symbol=symbol_name, signature=signature, purpose=override["purpose"], reason=override["reason"])
    return FunctionDetail(
        symbol=symbol_name,
        signature=signature,
        purpose=infer_function_purpose(relative_path, symbol_name),
        reason=infer_function_reason(relative_path, symbol_name),
    )


def extract_python_function_details(content: str, relative_path: str) -> list[FunctionDetail]:
    details: list[FunctionDetail] = []
    try:
        tree = ast.parse(content)
    except Exception:
        return details

    for node in tree.body:
        if isinstance(node, ast.FunctionDef):
            signature = format_python_signature(node)
            details.append(build_function_detail(relative_path, node.name, signature))
        elif isinstance(node, ast.ClassDef):
            for child in node.body:
                if isinstance(child, ast.FunctionDef):
                    full_name = f"{node.name}.{child.name}"
                    signature = format_python_signature(child)
                    details.append(build_function_detail(relative_path, full_name, signature))
    return details


def extract_ts_function_details(content: str, relative_path: str) -> list[FunctionDetail]:
    details: list[FunctionDetail] = []
    patterns = [
        re.compile(r"^(?:export\s+)?function\s+([A-Za-z_][A-Za-z0-9_]*)\s*(\([^)]*\))"),
        re.compile(r"^(?:export\s+)?const\s+([A-Za-z_][A-Za-z0-9_]*)\s*=\s*((?:async\s*)?\([^)]*\)\s*=>)"),
    ]
    for line in content.splitlines():
        raw = line.rstrip()
        if not raw or raw.startswith((" ", "\t")):
            continue
        for pattern in patterns:
            matched = pattern.search(raw)
            if matched:
                name = matched.group(1)
                signature = matched.group(2).replace("=>", "").strip()
                if not signature.startswith("("):
                    signature = "()"
                details.append(build_function_detail(relative_path, name, signature))
                break
    return details


def detect_tech_keywords(content: str, relative_path: str) -> str:
    tags: list[str] = []
    lower = content.lower()
    if "django" in lower:
        tags.append("Django/DRF")
    if "simplejwt" in lower or "jwt" in lower:
        tags.append("JWT 鉴权")
    if "chromadb" in lower:
        tags.append("Chroma 向量检索")
    if "from pptx" in lower or "python-pptx" in lower:
        tags.append("PPT 解析/编辑")
    if "fitz" in lower or "pymupdf" in lower:
        tags.append("PDF 解析/导出")
    if "from pil" in lower or "pillow" in lower:
        tags.append("图像处理")
    if "threadpoolexecutor" in lower or "threading" in lower:
        tags.append("并发/后台任务")
    if "requests" in lower and "chat" in lower:
        tags.append("大模型接口调用")
    if "axios" in lower:
        tags.append("Axios HTTP")
    if "markdown-it" in lower:
        tags.append("Markdown 渲染")
    if "vue-router" in lower:
        tags.append("前端路由")
    if "element-plus" in lower:
        tags.append("Element Plus UI")
    if "from docx" in lower:
        tags.append("Word 文档生成")
    if "tkinter" in lower:
        tags.append("桌面启动器")
    if relative_path.endswith("requirements.txt"):
        tags.append("Python 依赖管理")
    if relative_path.endswith("package.json"):
        tags.append("前端依赖与脚本")
    if not tags:
        tags.append("工程配置/通用脚本")
    return "、".join(dict.fromkeys(tags))


def infer_default_role(relative_path: str, ext: str) -> str:
    rp = relative_path.replace("\\", "/")
    name = Path(rp).name
    if "migrations/" in rp:
        return "该文件用于记录数据库模型演进过程，保证不同环境下数据结构可以按版本顺序升级。"
    if "/tests/" in rp:
        return "该文件用于自动化验证接口或服务行为，确保关键链路在修改后仍保持稳定。"
    if rp.endswith("/urls.py"):
        return "该文件负责定义路由映射关系，把外部访问路径连接到具体处理逻辑。"
    if rp.endswith("/settings.py"):
        return "该文件负责集中管理项目全局配置，包括数据库、应用注册、中间件和环境变量读取。"
    if rp.endswith("/models.py"):
        return "该文件负责定义持久化数据结构，是业务对象与数据库表之间的映射中心。"
    if rp.endswith("/serializers.py"):
        return "该文件负责请求参数和响应结构校验，是接口边界的重要保护层。"
    if rp.endswith("/views.py"):
        return "该文件负责接口编排与请求响应处理，将前端动作连接到后端服务逻辑。"
    if rp.endswith(".vue"):
        return "该文件是前端页面或布局组件，负责组织界面结构、状态和交互行为。"
    if rp.endswith((".ts", ".js")):
        return "该文件是前端脚本或配置文件，负责提供路由、类型、接口或共享逻辑。"
    if rp.endswith(".py") or rp.endswith(".pyw"):
        return "该文件是 Python 业务或工具脚本，承担某个明确的处理职责。"
    if name in {".env", ".env.example"}:
        return "该文件用于维护本地环境变量与服务运行参数。"
    return "该文件为项目中的辅助资源或配置产物。"


def infer_default_hard(content: str, relative_path: str, line_count: int) -> str:
    lower = content.lower()
    if "threadpoolexecutor" in lower or "threading" in lower:
        return "该文件涉及并发或后台流程控制，难点在于状态一致性、异常隔离和任务结束条件。"
    if "chromadb" in lower:
        return "该文件涉及向量索引构建与检索，难点在于召回质量与索引一致性。"
    if "pythoncom" in lower or "win32com" in lower:
        return "该文件依赖 Office COM 自动化，难点在于环境兼容性与进程稳定性。"
    if "json" in lower and "chat" in lower:
        return "该文件与模型结构化输出有关，难点在于容错解析与兜底策略设计。"
    if relative_path.endswith(".vue") and line_count > 400:
        return "该页面交互状态较多，难点在于控制副作用和避免不同业务链路互相干扰。"
    if "/tests/" in relative_path.replace("\\", "/"):
        return "测试文件的难点在于覆盖真实边界，同时降低对外部依赖波动的敏感度。"
    if line_count > 300:
        return "该文件职责较重，阅读时需要特别关注子函数拆分和模块边界设计。"
    return "该文件实现难度中等，主要考验职责拆分是否清晰和上下游接口是否稳定。"


def build_symbol_summary(content: str, path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".py", ".pyw"}:
        classes, funcs, imports = extract_python_symbols(content)
        parts: list[str] = []
        if classes:
            parts.append("类：" + "、".join(classes[:8]))
        if funcs:
            parts.append("函数：" + "、".join(funcs[:10]))
        if imports:
            parts.append("关键依赖：" + "、".join(imports[:8]))
        return "；".join(parts) if parts else "无显式类或函数定义。"
    if ext in {".ts", ".js", ".vue"}:
        symbols = extract_ts_like_symbols(content)
        return "主要函数：" + "、".join(symbols) if symbols else "以状态声明、模板和配置为主。"
    return "配置或静态内容文件。"


def build_function_details(content: str, relative_path: str, path: Path) -> list[FunctionDetail]:
    ext = path.suffix.lower()
    if ext in {".py", ".pyw"}:
        return extract_python_function_details(content, relative_path)
    if ext in {".ts", ".js", ".vue"}:
        return extract_ts_function_details(content, relative_path)
    return []


def build_file_record(path: Path) -> dict[str, Any]:
    relative = path.relative_to(ROOT).as_posix()
    is_text = is_text_file(path)
    content = safe_read_text(path) if is_text else ""
    lines = content.splitlines() if is_text else []
    line_count = len(lines)
    nonempty = len([line for line in lines if line.strip()]) if is_text else 0
    override = FILE_NOTE_OVERRIDES.get(relative, {})
    role = override.get("role") or infer_default_role(relative, path.suffix.lower())
    tech = override.get("tech") or (detect_tech_keywords(content, relative) if is_text else "文档/资源产物")
    hard = override.get("hard") or (infer_default_hard(content, relative, line_count) if is_text else "主要作为资源产物使用。")
    return {
        "path": relative,
        "is_text": is_text,
        "line_count": line_count,
        "nonempty": nonempty,
        "role": role,
        "tech": tech,
        "hard": hard,
        "symbols": build_symbol_summary(content, path) if is_text else "二进制或资源文件，不适合做源码级符号提取。",
        "functions": build_function_details(content, relative, path) if is_text else [],
        "size": path.stat().st_size,
    }


def configure_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)


def add_paragraph(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.bold = bold


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)


def add_function_block(doc: Document, detail: FunctionDetail) -> None:
    add_paragraph(doc, f"- {detail.symbol}{detail.signature}")
    add_paragraph(doc, f"  功能说明：{detail.purpose}")
    add_paragraph(doc, f"  这样设计的原因：{detail.reason}")


def generate_document(records: list[dict[str, Any]]) -> None:
    doc = Document()
    configure_style(doc)
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    add_title(doc, "双语课程辅助学习平台技术说明文档")
    add_paragraph(doc, f"生成时间：{now_str}")
    add_paragraph(doc, "说明：本文档基于当前工程源码自动生成，重点强化“文件作用说明”和“函数设计说明”。")

    text_records = [record for record in records if record["is_text"]]
    binary_records = [record for record in records if not record["is_text"]]
    total_lines = sum(record["line_count"] for record in text_records)
    total_nonempty = sum(record["nonempty"] for record in text_records)
    total_functions = sum(len(record["functions"]) for record in text_records)

    add_heading(doc, "1. 项目总体技术说明", 1)
    add_paragraph(
        doc,
        "本项目采用 Django + Vue3 的前后端分离架构，以大模型能力为中心，围绕课件学习场景形成了上传解析、双语翻译、译后预览、RAG 问答、总结生成、历史复盘和导出下载的完整闭环。"
    )
    add_paragraph(
        doc,
        "后端采用服务层拆分思路，将解析、翻译、导出、检索、总结、模型调用等高复杂度逻辑收敛到独立模块；前端则以工作台与历史页为核心，组织多状态页面交互。整体设计强调三点：一是结构化数据在链路中的贯通，二是长耗时任务的可观察性，三是模型不稳定输出下的兜底能力。"
    )

    add_heading(doc, "2. 关键设计决策", 1)
    for point in [
        "统一布局模型：无论源文件来自 PPT 还是 PDF，最终都转换为包含 page_width、page_height、text_containers、blocks 的统一结构，降低后续模块耦合。",
        "结构化翻译优先：翻译器优先让模型按容器级 JSON 返回结果，而不是只返纯文本，以便最大程度保留版式边界。",
        "缓存与并发协同：通过缓存减少重复翻译，再通过分块并发提升整份课件翻译速度，兼顾质量与性能。",
        "导出链路按格式分开处理：PPT 内部有可编辑文本框，所以系统直接把译文写回原始文本对象；PDF 更像定版页面，不适合稳定修改原文字对象，所以系统先按页面坐标重新绘制译文，再重新生成新的 PDF 文件。",
        "前端轮询而非阻塞等待：翻译任务启动后由前端持续刷新状态，保证长任务可观察且用户界面不中断。",
        "模型输出强约束 + 兜底：问答、总结、翻译都尽量要求结构化输出，同时准备本地回退逻辑，避免页面空白。"
    ]:
        add_paragraph(doc, f"- {point}")

    add_heading(doc, "3. 工程规模统计", 1)
    add_paragraph(doc, f"- 文本源码文件数：{len(text_records)}")
    add_paragraph(doc, f"- 文本源码总行数（含空行）：{total_lines}")
    add_paragraph(doc, f"- 文本源码总行数（非空行）：{total_nonempty}")
    add_paragraph(doc, f"- 识别到的主要函数/方法数量：{total_functions}")
    add_paragraph(doc, f"- 非源码资源文件数：{len(binary_records)}")

    add_heading(doc, "4. 逐文件与函数详解", 1)
    grouped: dict[str, list[dict[str, Any]]] = {}
    for record in records:
        top = record["path"].split("/")[0]
        grouped.setdefault(top, []).append(record)

    for group_index, top in enumerate(sorted(grouped.keys()), start=1):
        add_heading(doc, f"4.{group_index} {top}", 2)
        for record in sorted(grouped[top], key=lambda item: item["path"]):
            add_heading(doc, record["path"], 3)
            if not record["is_text"]:
                add_paragraph(doc, "作用说明：该文件属于资源、媒体、日志或其他非源码产物，不参与函数级设计说明。")
                add_paragraph(doc, f"文件大小：{record['size']} bytes")
                continue

            add_paragraph(doc, f"作用说明：{record['role']}")
            add_paragraph(doc, f"涉及技术/算法：{record['tech']}")
            add_paragraph(doc, f"设计难点：{record['hard']}")
            add_paragraph(doc, f"关键符号概览：{record['symbols']}")
            add_paragraph(doc, f"代码规模：{record['line_count']} 行（非空 {record['nonempty']} 行）")

            if record["functions"]:
                add_paragraph(doc, "函数设计说明：")
                for detail in record["functions"]:
                    add_function_block(doc, detail)
            else:
                add_paragraph(doc, "函数设计说明：该文件以配置、样式、模板或静态数据为主，没有需要单独展开的函数设计。")

    add_heading(doc, "5. 结论", 1)
    add_paragraph(
        doc,
        "从工程实现上看，该项目已经形成较清晰的分层：解析层负责提取可翻译结构，翻译层负责语言转换与容错，导出层负责视觉回写，检索与总结层负责学习辅助能力，前端负责将长链路过程变成可交互的工作台。"
    )
    add_paragraph(
        doc,
        "本次文档重写的重点，是把“文件为什么存在”和“函数为什么要这样拆分”明确写清楚，便于后续答辩、代码讲解和团队交接时直接使用。"
    )

    doc.save(str(OUT_PATH))


def main() -> None:
    files: list[Path] = []
    for path in ROOT.rglob("*"):
        if not path.is_file():
            continue
        if is_excluded(path.relative_to(ROOT)):
            continue
        if should_skip_file(path):
            continue
        files.append(path)
    records = [build_file_record(path) for path in sorted(files)]
    generate_document(records)
    print(str(OUT_PATH))


if __name__ == "__main__":
    main()
